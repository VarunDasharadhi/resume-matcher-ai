"""Resume text extraction for supported document types.

Supports PDF (via PyMuPDF), modern Word ``.docx`` (via python-docx), and a
best-effort reader for legacy Word ``.doc`` (via olefile). The ``.doc`` reader
pulls readable text out of the binary WordDocument stream, which is good enough
for skill matching; if it cannot recover usable text it raises so the caller
can ask the user to upload a PDF or .docx instead.
"""
import os
import re
from io import BytesIO

import fitz  # PyMuPDF
import olefile
from docx import Document

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}

# A .doc that yields fewer than this many real words is treated as unreadable.
_MIN_DOC_WORDS = 20
_WORD_RE = re.compile(r"[A-Za-z]{2,}")
# Keep tabs/newlines, printable ASCII and printable Unicode; drop control bytes.
_NON_PRINTABLE = re.compile("[^\t\n\r\x20-\x7e -￿]+")


def _ext(path: str) -> str:
    return os.path.splitext(path)[1].lower()


def is_supported(filename: str) -> bool:
    """True if ``filename``'s extension is one we can extract text from."""
    return _ext(filename or "") in SUPPORTED_EXTENSIONS


def extract_text_from_pdf(source) -> str:
    """Extract text from a PDF given a filesystem path or raw bytes."""
    if isinstance(source, (bytes, bytearray)):
        doc = fitz.open(stream=source, filetype="pdf")
    else:
        doc = fitz.open(source)
    try:
        return "".join(page.get_text() for page in doc)
    finally:
        doc.close()


def extract_text_from_docx(docx_path: str) -> str:
    document = Document(docx_path)
    parts = [p.text for p in document.paragraphs]
    # Include text from tables, which resumes often use for layout.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _clean(text: str) -> str:
    text = _NON_PRINTABLE.sub(" ", text)
    return re.sub(r"[ \t]{2,}", " ", text).strip()


def _word_count(text: str) -> int:
    return len(_WORD_RE.findall(text))


def _decode_doc_bytes(data: bytes) -> str:
    """Decode raw .doc stream bytes, picking the encoding that reads best.

    Word stores text as either UTF-16LE or an 8-bit codepage depending on the
    document, so we try a few and keep whichever recovers the most real words.
    """
    best, best_score = "", -1
    for enc in ("utf-16-le", "cp1252", "latin-1"):
        try:
            candidate = _clean(data.decode(enc, errors="ignore"))
        except Exception:
            continue
        score = _word_count(candidate)
        if score > best_score:
            best, best_score = candidate, score
    return best


def extract_text_from_doc(source) -> str:
    """Best-effort text extraction from a legacy ``.doc`` (path or file-like)."""
    if not olefile.isOleFile(source):
        raise ValueError("Not a valid .doc file")
    if hasattr(source, "seek"):
        source.seek(0)  # isOleFile consumed the header
    ole = olefile.OleFileIO(source)
    try:
        if not ole.exists("WordDocument"):
            raise ValueError("No WordDocument stream in .doc")
        data = ole.openstream("WordDocument").read()
    finally:
        ole.close()

    text = _decode_doc_bytes(data)
    if _word_count(text) < _MIN_DOC_WORDS:
        raise ValueError("Could not extract readable text from .doc")
    return text


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Extract text from an in-memory document, dispatching by extension.

    Nothing touches disk, which keeps the request path fast and fully
    stateless. Raises ``ValueError`` for unsupported file types.
    """
    ext = _ext(filename)
    if ext == ".pdf":
        return extract_text_from_pdf(data)
    if ext == ".docx":
        return extract_text_from_docx(BytesIO(data))
    if ext == ".doc":
        return extract_text_from_doc(BytesIO(data))
    raise ValueError(f"Unsupported file type: {ext or '(none)'}")


def extract_text(path: str) -> str:
    """Extract text from a supported document on disk.

    Raises ``ValueError`` for unsupported file types.
    """
    with open(path, "rb") as fh:
        data = fh.read()
    return extract_text_from_bytes(data, path)
